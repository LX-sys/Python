import os.path
import threading
import configparser
import time
import docx
from docx.shared import RGBColor
from docx.shared import Pt
# import platform
import re
# import sys
# import os
# __path="\\".join(os.getcwd().split("\\"))
# __path+="\\"+"run"+"\\"+"Keyword"
# print(__path)
# sys.path.append(__path)
# from Keyword import *
# vscode下导入其它文件夹下的模块的方法
#  ----
from Keyword.Keyword import Keyword
from PErrorcolor.PErrorcolor import pec


# 管理文件路径类
class AdminFilepath:
    HEADINFO = ""

    def __init__(self):
        self.doc = docx.Document()
        # 路径恢复标识
        self.__pathChanage = False
        # 临时路径
        self.__temp_path = []
        # 存放当个或多个文件路径,
        self.__path = []
        # 递归文件存放列表
        self.__recursionFile = []
        # 文件名列表
        self.__filename = []
        # 每个级别错误的默认颜色
        self.__color = Keyword.DEFAUKT_COLOR
        # 每个级别错误的默认字体大小
        self.__size = Keyword.DEFAUKT_SIZE
        # typora语法格式
        self.__typora = "<font color@c size@s>text</font>"
        # 文件内容映射表(保留语句颜色状态)
        self.__fileContent_dict = {}
        # 配置
        self.init_config()

    # 配置文件初始化
    def init_config(self):
        # 清除之前残留的映射
        self.__fileContent_dict.clear()
        # 创建配置文件对象
        self.config = configparser.ConfigParser()
        # 读取配置文件
        # 注意编码格式
        try:
            # 注意这里的路径问题,这个路径是相对于main.py文件来找的config.ini配置文件
            # 如果只执行当前文件路径改为 ..\config\config.ini
            p = self.config.read(r"config\config.ini", encoding="utf-8")
        except Exception:
            p = self.config.read(r"config\config.ini", encoding="gbk")
        # 当配置文件不存在,紧急配置
        if not p:

            pec.addtext({"err": ["1000", "配置文件config.ini 以丢失或者路径不对"]})
            # 是否需要 文件内容映射表(保留语句颜色状态) True/False
            self.__fileContent_state = False
        else:
            try:
                # 是否启动 文件内容映射表
                self.__fileContent_state = self.config["state"]["fileContent_state"]  # 默认启动
                if self.__fileContent_state == "True":
                    self.__fileContent_state = True
                elif self.__fileContent_state == "False":
                    self.__fileContent_state = False
                # 是否启动 头信息
                self.__headinfo_state = self.config["state"]["headinfo"]
                if self.__headinfo_state == "True":
                    self.__headinfo_state = True
                elif self.__headinfo_state == "False":
                    self.__headinfo_state = False
                # 获取自定义错误级别关键字
                if self.config["errorLevel"]["custom"]:
                    for lev in self.config["errorLevel"]["custom"].split(","):
                        if lev not in Keyword.KEYWORDS_LIST:
                            Keyword.ERRORLEVEL_LIST.append(lev)  # 加入错误关键字列表
                            Keyword.KEYWORDS_LIST.append(lev)  # 加入所有关键字集合
            except Exception:
                pec.addtext({"err": ["1001", "配置文件config.ini内容错误"]})
                self.__fileContent_state = False
                self.__headinfo_state = False

    # 恢复默认配置文件
    def config_default(self):
        # 配置文件内容
        data_dict = {"state": {"#是否需要 文件内容映射表(保留语句颜色状态)True/False": "",
                               "fileContent_state": "True",
                               "headinfo": "False"},
                     "errorLevel": {"# 自定义可扩展的错误级别 eg: custom = de,crl,op": "",
                                    "custom": ""}
                     }
        # 创建配置具体内容
        self.config.read_dict(data_dict)
        # 新建config.ini配置文件
        configFile = os.getcwd() + "\\" + "config" + "\\" + "config.ini"
        try:
            with open(configFile, "w") as f:
                self.config.write(f)
        except Exception:
            pec.addtext({"err": ["1002", "run目录下缺少config文件夹,请新建"]})
            return None

    # 设置颜色的接口
    def setColor(self, level: dict):
        if type(level) == dict:
            self.__color = level
        else:
            raise ("参数是类型是dict")

    # 获取时间
    def __getTime(self):
        # 获取时间
        mytime = "**" + "时间/time: " + time.strftime("%Y-%m-%d %H:%M:%S", time.localtime()) + "**" + "\n"
        return mytime

    # 头信息
    def up_headinfo(self):
        if self.__headinfo_state:
            mytime = self.__getTime()
            AdminFilepath.HEADINFO = mytime + self.__getheadinfo()

    # 设置头信息
    def setheadinfo(self, headinfo: str):
        if headinfo:
            self.__headinfo = headinfo
        else:
            self.__headinfo = ""
        if not self.__headinfo_state:
            pec.addtext({"err": ["1003", "使用此功能请先在配置文件中启动config.ini"]})

    # 获取头信息
    def __getheadinfo(self):
        try:
            return self.__headinfo
        except Exception:
            return ""

    # 获取颜色接口
    def getColor(self):
        return self.__color

    # 通过颜色获取类型
    def colorTotype(self, color: str):
        for ty in Keyword.DEFAUKT_COLOR_:
            if Keyword.DEFAUKT_COLOR_[ty] == color:
                return ty
        return None

    # 检测路径是否为空
    def isPathNone(self):
        if not self.__path:
            pec.addtext({"err": ["1004", "请先连接文件在执行命令 connect filePath1, filePath2, ..."]})
            return True
        else:
            return False

    # 重置路径
    def resetPath(self, path):
        temp = self.__path[::]  # 做一次深拷贝
        self.__filename.clear()  # 清除原来的文件名
        self.clear()  # 清除原来的路径
        self.setpathAndfileName(path)
        if self.__pathChanage:
            # 如更改后的路径有误，则更改无效
            self.__path = temp
            self.__pathChanage = False

    # 接收路径
    def __setPath(self, *args):
        # 接收所有输入,形成一个列表
        if args:
            for filepath in args:
                # 判断路径是否存在
                if filepath not in self.__path:
                    self.__temp_path.append(filepath)  # 先存入临时路径，等待处理
            # 解决在win下复制路径会导致路径多出'\u202a'问题,获取其它字符
            for i in range(len(self.__temp_path)):
                if len(self.__temp_path[i].split(":")[0]) != 1:
                    path = self.__temp_path[i].split(":")[0][-1] + ":" + self.__temp_path[i].split(":")[-1]
                    self.__temp_path[i] = path
            #     if "\u202a" in self.__temp_path[i]:
            #         self.__temp_path[i] = self.__temp_path[i].replace("\u202a", "")

    # 判断全是文件类型是否相同
    def __isfileTypeSome(self):
        # 路径列表不为None
        if self.__temp_path:
            # 判断是否为全是文件路径或者全是文件夹路径
            files_list = [os.path.isfile(file) for file in self.__temp_path]
            dirs_list = [os.path.isdir(file) for file in self.__temp_path]
            if files_list.count(True) == len(files_list) or dirs_list.count(True) == len(dirs_list):
                return True
            else:
                return False

    # 获取文件的类型(文件/文件夹)
    def __getfileType(self):
        files_list = [os.path.isfile(file) for file in self.__temp_path]
        dirs_list = [os.path.isdir(file) for file in self.__temp_path]
        if files_list.count(True) == len(files_list):
            return "file"
        if dirs_list.count(True) == len(dirs_list):
            return "files"

    # 设置路径，分离出文件名
    def setpathAndfileName(self, path_list):
        # 设置路径
        self.__setPath(*path_list)
        # 先判断文件类型是否一至
        if self.__isfileTypeSome():
            # 再判断是否未文件夹，文件不用判断，直接分离
            if self.__getfileType() == "files":
                # 遍历列表中所有文件夹
                for p in self.__temp_path:
                    self.recursionDIR(p)
                self.__temp_path = self.__recursionFile
            # 分离出文件名
            for name in self.__temp_path:
                temp_name = name.split("\\")[-1]
                # 文件是否存在
                if temp_name not in self.__filename:
                    self.__filename.append(temp_name)
            # 再次检测路径是否有重复
            for file in self.__temp_path[::]:  # 深拷贝
                if file not in self.__path:
                    self.__path.append(file)
        else:
            pec.addtext({"err": ["1005", "路径类型不一致!"]})
            self.__pathChanage = True  # 路径恢复
        # 清楚临时路径
        self.__temp_path.clear()
        self.__recursionFile.clear()

    # show路径+文件名
    def getfileName(self):
        # 路径列表不为None
        if self.__path:
            # 打印当前连接的文件
            for name in self.__path:
                print(name)
        else:
            print("Empty")

    # show config信息(不包括配置文件中的信息)
    def getconfig(self):
        print("[-----------------]")
        print("当前级别-->颜色")
        for k, v in self.__color.items():
            print(k, "-->", v)
        print("当前字体-->大小")
        for k, v in self.__size.items():
            print(k, "-->", v)
        print("[-----------------]")

    # docx写入加颜色"#698B22"
    def docxWriteColor(self, data, color, filename, ver=True):
        if color[0] == "#" and len(color) == 7:
            color = color[1:]  # 去除#
            # 每两个一组分割
            s, m, e = color[:2], color[2:4], color[4:6]
            r = int(Keyword.BIN_DICT[s[0]] + Keyword.BIN_DICT[s[1]], 2)
            g = int(Keyword.BIN_DICT[m[0]] + Keyword.BIN_DICT[m[1]], 2)
            b = int(Keyword.BIN_DICT[e[0]] + Keyword.BIN_DICT[e[1]], 2)
        else:
            pec.addtext({"err": ["1010", "请转入正确的颜色值"]})
            return
        pa = self.doc.add_paragraph()
        if data and ver:
            run = pa.add_run(data)
            run.font.size = Pt(12)  # word字体大小
            run.font.color.rgb = RGBColor(r, g, b)  # 颜色
        else:
            run = pa.add_run(data)
            run.font.size = Pt(12)
        self.doc.save(filename)

    # docx写入加字体
    def docxWriteSize(self, data, size, filename,ver=True):
        if int(size) <= 0:
            pec.addtext({"err": ["1011","字体大小不能为0"]})
            return
        pa = self.doc.add_paragraph()
        if data and ver == True:
            run = pa.add_run(data)
            run.font.size = Pt(int(size))
        else:
            run = pa.add_run(data)
            run.font.size = Pt(12)
        self.doc.save(filename)

    # 读取文件(每次读取一行,直到读取完)
    def myReadLine(self, filepath: str, errorKeyword, colorsize, cs="#698B22"):
        if  cs[0] == "#" and len(cs) == 7:
            pass
        else:
            if colorsize == "color":
                pec.addtext({"err": ["1010", "请转入正确的颜色值"]})
                return
        #  先检测文件编码格式
        try:
            with open(filepath, "r", encoding='gbk') as f:
                pass
            encoding = "gbk"
        except Exception:
            encoding = "UTF-8"
        # 获取文件名
        temp_file_name = filepath.split("\\")[-1]
        # 取出后缀
        suffix = temp_file_name.split(".")[-1]
        # 读文件
        if suffix == "md":
            with open(filepath, "r", encoding=encoding) as f:
                # 循环读取暑假，并存入列表
                line_list = []
                temp = []
                while True:
                    line = f.readline()
                    if line:
                        if line != "\n":  # 去除空白行
                            line_list.append(line)
                    else:
                        break
                count_line = 0  # 记录行数
                for text in line_list:
                    if re.findall(errorKeyword, text, flags=re.I):
                        if colorsize == Keyword.COLOR:  # 对文字颜色进行设置
                            Sentence = self.__typora.replace("@c", "=" + cs)
                            Sentence = Sentence.replace("@s", "")
                        elif colorsize == Keyword.SIZE:  # 对字体大小进行设置
                            Sentence = self.__typora.replace("@s", "=" + cs)
                            Sentence = Sentence.replace("@c", "")
                        Sentence = Sentence.replace("text", "**" + text.replace("\n", "") + "**")
                        Sentence += "\n"
                        temp.append(Sentence)  # 写入变化的语句
                        if self.__fileContent_state:
                            if temp_file_name not in self.__fileContent_dict:
                                self.__fileContent_dict[temp_file_name] = {}
                            # 标记行数所对应的语句
                            self.__fileContent_dict[temp_file_name][str(count_line)] = Sentence
                    else:
                        # 语句状态是否需要保留
                        if temp_file_name in self.__fileContent_dict and self.__fileContent_state:
                            if str(count_line) in self.__fileContent_dict[temp_file_name]:
                                temp.append(self.__fileContent_dict[temp_file_name][str(count_line)])
                            else:
                                temp.append(text)
                        else:
                            temp.append(text)
                    count_line += 1
        elif suffix in ["doc", "docx"]:
            # 读取wrod文档
            line_list = []
            doc = docx.Document(filepath)
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    line_list.append(paragraph.text)  # 加入内容列表

        # 插入头信息
        if self.__headinfo_state:
            temp.insert(0, AdminFilepath.HEADINFO)
        if suffix == "md":
            # 写文件(根据文件名)
            filename = temp_file_name
            if suffix in Keyword.FORMAT_LIST:  # 检查格式是否支持
                filename = filename.split(".")[0] + "_ECC" + "." + suffix  # ECC (ErrorColorConifg)
                # 判断当前文件是否在映射表中
                with open(filename, "w") as f:
                    f.write("".join(temp))  # 将列表连接成字符串，写入文件
            else:
                pec.addtext({"err": ["1009", "不支持当前后缀,支持{}".format(Keyword.FORMAT_LIST)]})
        elif suffix in ["doc", "docx"]:
            # 对word文档颜色/大小设置
            for text in line_list:
                print("re ",re.findall(errorKeyword, text, flags=re.I))
                if re.findall(errorKeyword, text, flags=re.I):
                    if colorsize == Keyword.COLOR:  # 对文字颜色进行设置
                        self.docxWriteColor(data=text, color=cs,
                                            filename=temp_file_name.split(".")[0] + "_ECC" + "." + suffix)
                    elif colorsize == Keyword.SIZE:  # 对字体大小进行设置
                        self.docxWriteSize(data=text,size=cs,
                                           filename=temp_file_name.split(".")[0] + "_ECC" + "." + suffix)
                else:
                    if colorsize == Keyword.COLOR:
                        self.docxWriteColor(data=text, color=cs,
                                            filename=temp_file_name.split(".")[0] + "_ECC" + "." + suffix, ver=False)
                    elif colorsize == Keyword.SIZE:
                        self.docxWriteSize(data=text, size=cs,
                                            filename=temp_file_name.split(".")[0] + "_ECC" + "." + suffix, ver=False)
    # 递归文件夹(存储)
    def recursionDIR(self, path):
        try:
            dir_list = os.listdir(path)
            for i in dir_list:
                sub_dir = os.path.join(path, i)
                if os.path.isdir(sub_dir):
                    self.recursionDIR(sub_dir)
                else:  # 此时sub_dir是文件的绝对路径
                    self.__recursionFile.append(sub_dir)
        except Exception:
            # 这里报错是因为文件列表中多出了一个目录，而且在开头
            self.__recursionFile.pop(0)

    # 递归文件夹(显示)(暂时不用此函数)
    def recursionDIRShow(self, path):
        dir_list = os.listdir(path)
        for i in dir_list:
            sub_dir = os.path.join(path, i)
            if os.path.isdir(sub_dir):
                self.recursionDIR(sub_dir)
            else:  # 此时sub_dir是文件的绝对路径
                print("  --->", sub_dir)
                self.__filename.append(sub_dir.split("\\")[-1])

    # 打开文件
    def opens(self, errorKeyword, colorsize, cs="#698B22"):
        if not self.__path:
            return False

        # 工作线程
        def work(filepath):
            print("---")
            self.myReadLine(filepath, errorKeyword, colorsize, cs)
            print("---")

        # 检查当前列表路径是否为文件夹
        if os.path.isdir(self.__path[0]):
            # 遍历列表中所有文件夹
            for p in self.__path:
                self.recursionDIR(p)
            self.__path = self.__recursionFile
        threadpool = []  # 线程池
        # 为每一个文件开一个线程
        for filepath in self.__path:
            print("open", filepath)
            threadpool.append(threading.Thread(
                target=work(filepath, )))  # 加入线程池
        # 启动线程,等待
        for s in threadpool:
            s.start()
            s.join()

    # 导出文件
    def out(self, level: str, path: str, format="md"):
        # 首先判断导出格式是否支持
        if format not in Keyword.FORMAT_LIST:
            pec.addtext({"err": ["1006", "不支持当前格式,支持:{}".format(Keyword.FORMAT_LIST)]})
            return
        def Rmd():
            text_list=[]
            with open(filepath, "r") as f:
                while True:
                    data = f.readline()
                    if not data:
                        break
                    if re.findall(level, data, flags=re.I):  # 不区分大小写
                        text_list.append(data)
            return text_list
        def Rword():
            text_list = []
            doc = docx.Document(filepath)
            for paragraph in doc.paragraphs:
                data = paragraph.text
                if data:
                    if re.findall(level, data, flags=re.I):  # 不区分大小写
                        text_list.append(data)  # 加入内容列表
            return text_list
        def Wmd(text_list,name,path):
            with open(path, "w") as f:
                f.write("".join(text_list))
            print(name, "....OK")
        def Wword(text_list,name,path):
            for text in text_list:
                self.docxWriteSize(data=text, size=1, filename=path, ver=False)
            print(name, "....OK")
        # 工作线程
        def work(filepath: str, path):
            # 取文件后缀
            tempFileSux=filepath.split("\\")[-1].split(".")[-1]

            # 读取文件
            if tempFileSux == "md":
                text_list=Rmd()
            elif tempFileSux in ["doc","docx"]:
                text_list=Rword()
            # 获取文件名称
            filename = filepath.split("\\")[-1].split(".")[0]
            name = level + "_" + filename + "_" + "_" + "副本" + "." + format
            path += "\\" + name
            # 插入头信息
            if self.__headinfo_state:
                text_list.insert(0, AdminFilepath.HEADINFO)
            print(format)
            # 写入文件
            if format == "md":
                Wmd(text_list, name,path)
            elif format in ["doc", "docx"]:
                Wword(text_list, name,path)

        threadpool = []  # 线程池
        # 为每一个文件开一个线程
        for filepath in self.__path:
            threadpool.append(threading.Thread(
                target=work(filepath, path, )))  # 加入线程池
        # 启动线程,等待
        for s in threadpool:
            s.start()
            s.join()

    # 清除路径
    def clear(self):
        self.__path.clear()


adp = AdminFilepath()
